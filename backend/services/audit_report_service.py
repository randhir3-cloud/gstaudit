"""Full GAIS audit report generation — preview and export."""

from __future__ import annotations

import io
from typing import List, Optional

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from models.audit_session import AuditSession
from models.investigation import InvestigationCase
from models.dealer_metadata import DealerMetadata
from services.comparison_store import get_result
from services.dashboard_service import build_dashboard
from services.investigation_service import build_summary, sync_cases_from_comparison
from services.report_export import safe_report_filename
from intelligence.intelligence_service import get_session_intelligence


def build_report_preview(session: AuditSession) -> dict:
    dash = build_dashboard(session)
    cmp_result = get_result(session.session_id)
    cases = sync_cases_from_comparison(session.session_id)
    case_summary = build_summary(cases)
    high_risk = [c for c in cases if c.risk_score >= 70]

    cmp_summary = cmp_result.summary.model_dump() if cmp_result else {}
    observations = []
    for i, case in enumerate(cases[:100], start=1):
        if case.result_type == "MATCHED":
            continue
        observations.append({
            "observation_no": i,
            "description": f"{case.result_type}: Invoice {case.invoice_number}",
            "applicable_dataset": "GSTR-1 / EWB OUTWARD",
            "risk": case.risk_score,
            "officer_remark": case.officer_remarks,
            "final_status": case.status,
        })

    datasets_uploaded = sum(1 for k, ds in (session.datasets or {}).items() if ds.merged or ds.source_files or ds.staged_files)
    risk_level = cmp_summary.get("risk_level", "LOW")
    if isinstance(risk_level, dict):
        risk_level = risk_level.get("value", "LOW")

    intel = get_session_intelligence(session.session_id) if cmp_result else None
    intel_summary = intel.summary.model_dump() if intel else None

    from services.msae_service import get_session_msae
    msae = get_session_msae(session.session_id) if cmp_result else None
    msae_report = None
    if msae and msae.summary.master_case_count > 0:
        msae_report = {
            "summary": msae.summary.model_dump(),
            "top_master_cases": [c.model_dump() for c in msae.master_cases[:10]],
            "patterns": [p.model_dump() for p in msae.patterns[:8]],
            "scores": msae.summary.scores.model_dump(),
        }

    return {
        "executive_summary": {
            "dealer": dash.dealer_name,
            "gstin": dash.gstin,
            "financial_year": dash.financial_year,
            "datasets_uploaded": datasets_uploaded,
            "comparison_performed": cmp_result is not None,
            "total_invoices": cmp_summary.get("total_gstr1_records", 0) + cmp_summary.get("total_eway_records", 0),
            "matched": cmp_summary.get("matched_count", 0),
            "missing": cmp_summary.get("missing_in_gstr1_count", 0) + cmp_summary.get("missing_in_eway_count", 0),
            "duplicates": cmp_summary.get("duplicate_count", 0),
            "risk_level": str(risk_level),
            "audit_conclusion": _audit_conclusion(case_summary, cmp_summary),
        },
        "dealer_profile": {
            "legal_name": session.dealer.legal_name,
            "trade_name": session.dealer.trade_name,
            "gstin": session.dealer.gstin,
            "financial_year": session.financial_year or session.dealer.financial_year,
            "tax_period": session.dealer.tax_period,
        },
        "audit_scope": {
            "comparison_pairs": ["GSTR-1 ↔ EWB OUTWARD"],
            "period": session.financial_year or session.dealer.financial_year,
        },
        "data_availability": {
            "dataset_cards": dash.dataset_cards,
            "readiness_percent": dash.audit_readiness_percent,
        },
        "comparison_summary": cmp_summary,
        "discrepancy_summary": dash.discrepancies.model_dump(),
        "case_tracking": case_summary.model_dump(),
        "high_risk_cases": [c.model_dump() for c in high_risk[:20]],
        "observations": observations,
        "recommendations": _recommendations(case_summary, cmp_summary, intel),
        "audit_intelligence": {
            "summary": intel_summary,
            "patterns": intel_summary.get("patterns", []) if intel_summary else [],
            "month_analysis": [m.model_dump() for m in intel.months[:12]] if intel else [],
            "top_suppliers": [s.model_dump() for s in intel.suppliers[:10]] if intel else [],
            "top_customers": [c.model_dump() for c in intel.customers[:10]] if intel else [],
            "high_risk_months": intel_summary.get("executive_insights", {}).get("months_requiring_verification", []) if intel_summary else [],
            "suggested_documents": [d.model_dump() for d in intel.document_recommendations[:8]] if intel else [],
            "priority_cases": intel_summary.get("priority_cases", []) if intel_summary else [],
        },
        "consolidated_audit": msae_report,
        "annexures": {
            "total_cases": len(cases),
            "export_note": "Full case annexure available via Excel export.",
        },
    }


def _audit_conclusion(case_summary, cmp_summary) -> str:
    if not cmp_summary:
        return "Audit data upload incomplete. Comparison not performed."
    missing = cmp_summary.get("missing_in_gstr1_count", 0) + cmp_summary.get("missing_in_eway_count", 0)
    if missing == 0 and case_summary.high_risk == 0:
        return "Records reconciled. No material discrepancies requiring further action."
    return f"Material discrepancies identified. {missing} missing records and {case_summary.high_risk} high-risk cases require officer verification."


def _recommendations(case_summary, cmp_summary, intel=None) -> List[str]:
    recs = []
    if cmp_summary.get("missing_in_gstr1_count"):
        recs.append("Verify GSTR-1 filing for invoices present in E-Way Bills only.")
    if cmp_summary.get("missing_in_eway_count"):
        recs.append("Verify E-Way Bill generation for invoices declared in GSTR-1 without corresponding EWB.")
    if cmp_summary.get("value_mismatch_count"):
        recs.append("Reconcile invoice value differences with books of account.")
    if case_summary.pending:
        recs.append(f"Resolve {case_summary.pending} pending investigation cases.")
    if intel and intel.summary.patterns:
        for p in intel.summary.patterns[:3]:
            recs.append(f"Pattern detected: {p.description}")
    if intel and intel.summary.executive_insights.months_requiring_verification:
        months = ", ".join(intel.summary.executive_insights.months_requiring_verification[:3])
        recs.append(f"Prioritize detailed verification for months: {months}.")
    if not recs:
        recs.append("Maintain current compliance controls. No immediate corrective action required.")
    return recs


def _section_rows(preview: dict) -> List[tuple[str, str]]:
    es = preview["executive_summary"]
    return [
        ("Dealer", es["dealer"]),
        ("GSTIN", es["gstin"]),
        ("Financial Year", es["financial_year"]),
        ("Datasets Uploaded", str(es["datasets_uploaded"])),
        ("Comparison Performed", "Yes" if es["comparison_performed"] else "No"),
        ("Total Invoices", str(es["total_invoices"])),
        ("Matched", str(es["matched"])),
        ("Missing", str(es["missing"])),
        ("Duplicates", str(es["duplicates"])),
        ("Risk Level", es["risk_level"]),
        ("Audit Conclusion", es["audit_conclusion"]),
    ]


def build_full_excel_report(session: AuditSession, cases: List[InvestigationCase]) -> io.BytesIO:
    preview = build_report_preview(session)
    wb = openpyxl.Workbook()

    def _write_sheet(ws, title: str, rows: List[tuple]):
        ws.title = title[:31]
        ws["A1"] = title
        ws["A1"].font = Font(bold=True, size=14, color="1E3A8A")
        r = 3
        for label, value in rows:
            ws.cell(row=r, column=1, value=label).font = Font(bold=True)
            ws.cell(row=r, column=2, value=str(value))
            r += 1
        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 50

    _write_sheet(wb.active, "Executive Summary", _section_rows(preview))

    ws2 = wb.create_sheet("Observations")
    ws2.append(["No.", "Description", "Dataset", "Risk", "Remark", "Status"])
    for obs in preview["observations"]:
        ws2.append([obs["observation_no"], obs["description"], obs["applicable_dataset"], obs["risk"], obs["officer_remark"], obs["final_status"]])

    ws3 = wb.create_sheet("Cases")
    ws3.append(["Case No.", "Invoice", "Type", "Risk", "Status", "Remarks", "Period"])
    for case in cases:
        ws3.append([case.case_number, case.invoice_number, case.result_type, case.risk_score, case.status, case.officer_remarks, case.source_period])

    ws4 = wb.create_sheet("Recommendations")
    for i, rec in enumerate(preview["recommendations"], start=1):
        ws4.cell(row=i, column=1, value=f"{i}. {rec}")

    ai = preview.get("audit_intelligence") or {}
    if ai:
        ws5 = wb.create_sheet("Audit Intelligence")
        ws5.append(["Pattern", "Description", "Severity", "Count"])
        for p in ai.get("patterns", [])[:30]:
            ws5.append([p.get("pattern_type", ""), p.get("description", ""), p.get("severity", ""), p.get("affected_count", 0)])
        ws5.append([])
        ws5.append(["Month", "Mismatch %", "Risk %", "Largest Diff"])
        for m in ai.get("month_analysis", [])[:12]:
            ws5.append([m.get("month", ""), m.get("mismatch_percent", 0), m.get("risk_percent", 0), m.get("largest_difference", 0)])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def build_full_pdf_report(session: AuditSession, cases: List[InvestigationCase]) -> io.BytesIO:
    preview = build_report_preview(session)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontSize=16, textColor=colors.HexColor("#1E3A8A"), spaceAfter=12)
    story = [Paragraph("GST Audit Intelligence Report", title_style), Paragraph("Executive Summary", styles["Heading2"]), Spacer(1, 0.2 * cm)]
    table_data = [["Field", "Value"]] + [[a, b] for a, b in _section_rows(preview)]
    table = Table(table_data, colWidths=[5.5 * cm, 11 * cm])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DBEAFE"))]))
    story.extend([table, Spacer(1, 0.5 * cm), Paragraph("Recommendations", styles["Heading2"])])
    for rec in preview["recommendations"]:
        story.append(Paragraph(f"• {rec}", styles["Normal"]))
    story.extend([Spacer(1, 0.5 * cm), Paragraph(f"Annexure: {len(cases)} investigation cases recorded.", styles["Normal"])])
    doc.build(story)
    buffer.seek(0)
    return buffer


def build_full_docx_report(session: AuditSession, cases: List[InvestigationCase]) -> io.BytesIO:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word export") from exc

    preview = build_report_preview(session)
    doc = Document()
    doc.add_heading("GST Audit Intelligence Report", 0)
    doc.add_heading("Executive Summary", level=1)
    for label, value in _section_rows(preview):
        doc.add_paragraph(f"{label}: {value}")
    doc.add_heading("Recommendations", level=1)
    for rec in preview["recommendations"]:
        doc.add_paragraph(rec, style="List Bullet")
    doc.add_heading("Observations", level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    for i, h in enumerate(["No.", "Description", "Dataset", "Risk", "Remark", "Status"]):
        hdr[i].text = h
    for obs in preview["observations"][:50]:
        row = table.add_row().cells
        row[0].text = str(obs["observation_no"])
        row[1].text = obs["description"]
        row[2].text = obs["applicable_dataset"]
        row[3].text = str(obs["risk"])
        row[4].text = obs["officer_remark"] or ""
        row[5].text = obs["final_status"]
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def report_filename(dealer: DealerMetadata, ext: str) -> str:
    return safe_report_filename(dealer, ext, prefix="GAIS_Audit_Report")
